# 📘 Study Scheduler Pro 

import pandas as pd
from datetime import datetime, timedelta
import tkinter as tk
from tkinter import ttk, messagebox
import random
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
import os

# === Funções principais ===

materias = {}
status_estudo = {}

def gerar_agendas():
    if not materias:
        messagebox.showwarning("Aviso", "Adicione pelo menos uma matéria antes de gerar a agenda.")
        return

    try:
        dias = int(entry_dias.get())
        horas_por_dia = float(entry_horas.get())
    except ValueError:
        messagebox.showerror("Erro", "Digite valores válidos para dias e horas.")
        return

    nome_usuario = entry_nome.get().strip() or "Estudante"
    inicio = datetime.now()
    total_pesos = sum(materias.values())
    hora_inicio_estudo = 8
    minutos_por_dia = int(horas_por_dia * 60)
    agenda = []

    for dia in range(dias):
        data_atual = inicio + timedelta(days=dia)
        hora_atual = hora_inicio_estudo * 60

        for materia, peso in materias.items():
            minutos_estudo = int((peso / total_pesos) * minutos_por_dia)
            hora_ini = hora_atual
            hora_fim = hora_atual + minutos_estudo
            hora_atual = hora_fim

            h_ini = f"{hora_ini // 60:02d}:{hora_ini % 60:02d}"
            h_fim = f"{hora_fim // 60:02d}:{hora_fim % 60:02d}"

            agenda.append({
                "Data": data_atual.strftime("%d/%m/%Y"),
                "Horário": f"{h_ini} - {h_fim}",
                "Matéria": materia,
                "Duração (h)": round(minutos_estudo / 60, 2)
            })

            pausa = random.randint(10, 15)
            hora_ini_pausa = hora_atual
            hora_fim_pausa = hora_atual + pausa
            hora_atual = hora_fim_pausa

            h_ini_p = f"{hora_ini_pausa // 60:02d}:{hora_ini_pausa % 60:02d}"
            h_fim_p = f"{hora_fim_pausa // 60:02d}:{hora_fim_pausa % 60:02d}"

            agenda.append({
                "Data": data_atual.strftime("%d/%m/%Y"),
                "Horário": f"{h_ini_p} - {h_fim_p}",
                "Matéria": "☕ Pausa",
                "Duração (h)": round(pausa / 60, 2)
            })

    df = pd.DataFrame(agenda)

    pasta_agendas = "Agendas_Geradas"
    os.makedirs(pasta_agendas, exist_ok=True)

    # === Criação do Documento ===
    doc = Document()

    # === Capa elegante ===
    doc.add_paragraph("\n\n\n\n")
    titulo = doc.add_heading("📘 AGENDA DE ESTUDOS", level=1)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo.runs[0].font.size = Pt(28)
    titulo.runs[0].font.color.rgb = RGBColor(0, 102, 204)

    doc.add_paragraph("\n")

    nome_paragrafo = doc.add_paragraph(f"Elaborado para: {nome_usuario}")
    nome_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    nome_paragrafo.runs[0].font.size = Pt(16)
    nome_paragrafo.runs[0].font.color.rgb = RGBColor(0, 51, 102)

    doc.add_paragraph("\n")

    data_paragrafo = doc.add_paragraph(f"Data de criação: {datetime.now().strftime('%d/%m/%Y')}")
    data_paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    data_paragrafo.runs[0].font.size = Pt(14)
    data_paragrafo.runs[0].font.color.rgb = RGBColor(80, 80, 80)

    doc.add_paragraph("\n\n")

    frases = [
        "“A disciplina é a ponte entre metas e conquistas.” – Jim Rohn",
        "“Não espere por oportunidades, crie-as.”",
        "“Estudar é plantar o futuro com as mãos do presente.”",
        "“Cada página estudada é um passo rumo ao seu sonho.”",
        "“Grandes conquistas começam com pequenos hábitos diários.”"
    ]
    frase = random.choice(frases)
    p_frase = doc.add_paragraph(frase)
    p_frase.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_frase = p_frase.runs[0]
    run_frase.font.size = Pt(13)
    run_frase.italic = True
    run_frase.font.color.rgb = RGBColor(30, 30, 30)

    doc.add_page_break()  # separa capa e agenda

    # === Corpo do documento ===
    subtitulo = doc.add_heading("📅 PLANO DE ESTUDOS DETALHADO", level=2)
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("\n")

    tabela = doc.add_table(rows=1, cols=4)
    tabela.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tabela.style = "Table Grid"

    cabecalhos = ["Data", "Horário", "Matéria", "Duração (h)"]
    hdr_cells = tabela.rows[0].cells
    for i, nome in enumerate(cabecalhos):
        p = hdr_cells[i].paragraphs[0]
        run = p.add_run(nome)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(255, 255, 255)
        hdr_cells[i]._element.get_or_add_tcPr().append(parse_xml(f'<w:shd {nsdecls("w")} w:fill="007ACC"/>'))
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    cores_materias = [
        "#D1E8FF", "#E0FFD1", "#FFEFD1", "#FFD1D1", "#E8D1FF",
        "#FFF7D1", "#D1FFF0", "#FFD1F7"
    ]
    cores_atribuídas = {}
    cor_index = 0

    def colorir_celula(cell, cor_hex):
        shading = parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), cor_hex.replace("#", "")))
        cell._element.get_or_add_tcPr().append(shading)

    for _, linha in df.iterrows():
        row_cells = tabela.add_row().cells
        materia = linha["Matéria"]

        if materia not in cores_atribuídas and materia != "☕ Pausa":
            cores_atribuídas[materia] = cores_materias[cor_index % len(cores_materias)]
            cor_index += 1

        cor_fundo = "#E6E6E6" if materia == "☕ Pausa" else cores_atribuídas[materia]

        for i, valor in enumerate(linha):
            p = row_cells[i].paragraphs[0]
            run = p.add_run(str(valor))
            run.font.size = Pt(11)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            colorir_celula(row_cells[i], cor_fundo)

    doc.add_paragraph("\n")
    rodape = doc.add_paragraph("Gerado automaticamente pelo Study Scheduler Pro ✨")
    rodape.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rodape.runs[0].italic = True
    rodape.runs[0].font.size = Pt(10)

    # === Salvamento ===
    agora = datetime.now().strftime("%Y-%m-%d_%H-%M")
    nome_arquivo = f"agenda_estudos_{agora}.docx"
    caminho_completo = os.path.join(pasta_agendas, nome_arquivo)
    doc.save(caminho_completo)

    messagebox.showinfo("Sucesso", f"✅ Agenda criada com sucesso!\n\nSalva em:\n{os.path.abspath(caminho_completo)}")
    os.startfile(pasta_agendas)

def adicionar_materia():
    nome = entry_materia.get().strip()
    try:
        peso = int(entry_peso.get())
    except ValueError:
        messagebox.showerror("Erro", "Digite um número válido para o peso.")
        return

    if not nome:
        messagebox.showwarning("Aviso", "Digite o nome da matéria.")
        return

    materias[nome] = peso
    atualizar_lista()
    entry_materia.delete(0, tk.END)
    entry_peso.delete(0, tk.END)

def remover_materia():
    selecionada = lista_materias.selection()
    if not selecionada:
        messagebox.showwarning("Aviso", "Selecione uma matéria para remover.")
        return

    nome = lista_materias.item(selecionada[0], "values")[0]
    materias.pop(nome, None)
    atualizar_lista()

def atualizar_lista():
    for i in lista_materias.get_children():
        lista_materias.delete(i)
    for nome, peso in materias.items():
        lista_materias.insert("", "end", values=(nome, peso))

def carregar_estudos_do_dia():
    lista_estudos.delete(*lista_estudos.get_children())
    for materia in materias.keys():
        lista_estudos.insert("", "end", values=(materia, "⏳ Pendente"))
        status_estudo[materia] = False

def marcar_como_concluida(event):
    item = lista_estudos.selection()
    if not item:
        return
    item_id = item[0]
    nome, status = lista_estudos.item(item_id, "values")

    if not status_estudo[nome]:
        lista_estudos.item(item_id, values=(nome, "✅ Concluída"))
        lista_estudos.item(item_id, tags=("done",))
        lista_estudos.tag_configure("done", background="#b6f0c1")
        status_estudo[nome] = True
    else:
        lista_estudos.item(item_id, values=(nome, "⏳ Pendente"))
        lista_estudos.item(item_id, tags=("pending",))
        lista_estudos.tag_configure("pending", background="white")
        status_estudo[nome] = False

# === INTERFACE ===
janela = tk.Tk()
janela.title("📘 Study Scheduler Pro — Modo Widescreen")
janela.geometry("1200x700")
janela.configure(bg="#f0f4fa")
janela.resizable(True, True)

style = ttk.Style(janela)
style.theme_use("clam")
style.configure("Treeview.Heading", background="#007ACC", foreground="white", font=("Arial", 11, "bold"))
style.configure("Treeview", background="white", fieldbackground="white", font=("Arial", 10))

# === Título superior ===
titulo = tk.Label(janela, text="📚 STUDY SCHEDULER PRO", bg="#f0f4fa",
                  fg="#003366", font=("Segoe UI", 22, "bold"))
titulo.pack(pady=10)

# === Divisão principal (widescreen) ===
frame_main = tk.Frame(janela, bg="#f0f4fa")
frame_main.pack(fill="both", expand=True, padx=15, pady=10)

# Painel esquerdo — Controle
frame_left = tk.Frame(frame_main, bg="#ffffff", relief="raised", bd=2)
frame_left.pack(side="left", fill="both", expand=True, padx=(0,10), pady=5)

# Painel direito — Agenda e progresso
frame_right = tk.Frame(frame_main, bg="#ffffff", relief="raised", bd=2)
frame_right.pack(side="right", fill="both", expand=True, padx=(10,0), pady=5)

# === PAINEL ESQUERDO ===
tk.Label(frame_left, text="💡 Configurações de Estudo", bg="#ffffff",
         fg="#003366", font=("Arial", 14, "bold")).pack(pady=10)

frame_nome = tk.Frame(frame_left, bg="#ffffff")
frame_nome.pack(pady=5)
tk.Label(frame_nome, text="Seu nome:", bg="#ffffff", font=("Arial", 11)).pack(side="left", padx=5)
entry_nome = tk.Entry(frame_nome, width=30)
entry_nome.pack(side="left", padx=5)

frame_materias = tk.Frame(frame_left, bg="#ffffff")
frame_materias.pack(pady=10)

tk.Label(frame_materias, text="Matéria:", bg="#ffffff", font=("Arial", 11)).grid(row=0, column=0, padx=5)
entry_materia = tk.Entry(frame_materias, width=22)
entry_materia.grid(row=0, column=1, padx=5)

tk.Label(frame_materias, text="Peso:", bg="#ffffff", font=("Arial", 11)).grid(row=0, column=2, padx=5)
entry_peso = tk.Entry(frame_materias, width=6)
entry_peso.grid(row=0, column=3, padx=5)

tk.Button(frame_materias, text="Adicionar", command=adicionar_materia,
          bg="#4CAF50", fg="white", font=("Arial", 10, "bold"), width=10).grid(row=0, column=4, padx=5)
tk.Button(frame_materias, text="Remover", command=remover_materia,
          bg="#F44336", fg="white", font=("Arial", 10, "bold"), width=10).grid(row=0, column=5, padx=5)

tk.Label(frame_left, text="📘 Matérias Cadastradas", bg="#ffffff",
         fg="#003366", font=("Arial", 13, "bold")).pack(pady=10)

lista_materias = ttk.Treeview(frame_left, columns=("Matéria", "Peso"), show="headings", height=10)
lista_materias.heading("Matéria", text="Matéria")
lista_materias.heading("Peso", text="Peso")
lista_materias.pack(padx=10, pady=10, fill="x")

frame_tempo = tk.Frame(frame_left, bg="#ffffff")
frame_tempo.pack(pady=10)
tk.Label(frame_tempo, text="Dias de estudo:", bg="#ffffff", font=("Arial", 11)).grid(row=0, column=0, padx=5)
entry_dias = tk.Entry(frame_tempo, width=5)
entry_dias.insert(0, "7")
entry_dias.grid(row=0, column=1, padx=5)

tk.Label(frame_tempo, text="Horas por dia:", bg="#ffffff", font=("Arial", 11)).grid(row=0, column=2, padx=5)
entry_horas = tk.Entry(frame_tempo, width=5)
entry_horas.insert(0, "4")
entry_horas.grid(row=0, column=3, padx=5)

# === PAINEL DIREITO ===
tk.Label(frame_right, text="📅 Progresso Diário", bg="#ffffff",
         fg="#003366", font=("Arial", 14, "bold")).pack(pady=10)

lista_estudos = ttk.Treeview(frame_right, columns=("Matéria", "Status"), show="headings", height=12)
lista_estudos.heading("Matéria", text="Matéria")
lista_estudos.heading("Status", text="Status")
lista_estudos.pack(padx=15, pady=10, fill="x")

lista_estudos.bind("<Double-1>", marcar_como_concluida)

tk.Button(frame_right, text="Carregar matérias do dia", command=carregar_estudos_do_dia,
          bg="#2196F3", fg="white", font=("Arial", 12, "bold"), width=30).pack(pady=10)

tk.Button(frame_right, text="📄 Gerar Agenda Word com Capa", command=gerar_agendas,
          bg="#007ACC", fg="white", font=("Arial", 12, "bold"), width=35).pack(pady=15)

footer = tk.Label(janela, text="Desenvolvido com ❤️ por Code GPT",
                  bg="#f0f4fa", fg="#444444", font=("Arial", 10, "italic"))
footer.pack(side="bottom", pady=5)


janela.mainloop()
